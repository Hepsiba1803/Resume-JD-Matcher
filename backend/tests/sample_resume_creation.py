from docx import Document
from fpdf import FPDF

# Create a sample DOCX file
def create_sample_docx(filename):
    doc = Document()
    doc.add_heading('Sample Resume', level=1)
    doc.add_paragraph('Name: John Doe')
    doc.add_paragraph('Email: john.doe@example.com')
    doc.add_paragraph('Experience: 5 years in software development')
    doc.add_paragraph('Skills: Python, FastAPI, Machine Learning')
    doc.save(filename)

# Create a sample PDF file
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Sample Resume', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_sample_pdf(filename):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    pdf.cell(0, 10, 'Name: John Doe', 0, 1)
    pdf.cell(0, 10, 'Email: john.doe@example.com', 0, 1)
    pdf.cell(0, 10, 'Experience: 5 years in software development', 0, 1)
    pdf.cell(0, 10, 'Skills: Python, FastAPI, Machine Learning', 0, 1)
    pdf.output(filename)

# Create the files
create_sample_docx('sample_resume.docx')
create_sample_pdf('sample_resume.pdf')
# The above code creates two sample files: a DOCX file and a PDF file.
# These files can be used for testing the resume upload functionality in the FastAPI application.