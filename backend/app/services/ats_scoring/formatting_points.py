import re
import pdfplumber
from docx import Document
from fastapi import UploadFile

def formatting_score_and_suggestions(resume: UploadFile) -> tuple:
    """
    Evaluates resume formatting for ATS-friendliness and offers suggestions in a friendly tone.
    Args:
        resume (UploadFile): The uploaded resume file.
    Returns:
        tuple: (score, feedback list)
    """
    feedback = []
    max_points = 15
    deductions = 0
    STANDARD_FONTS = {'arial', 'calibri', 'times new roman'}
    DATE_PATTERNS = [
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}",
        r"\d{2}/\d{4}",
        r"\d{4}"
    ]
    SIMPLE_FILENAME_RE = re.compile(r"^[a-zA-Z0-9_\-]+(\.pdf|\.docx)$")

    resume.file.seek(0)
    resume_name = resume.filename

    # ----------- FILE NAME CHECK -----------
    if not SIMPLE_FILENAME_RE.match(resume_name):
        deductions += 1
        feedback.append("• Rename your file with only alphabets, numbers, underscores, or dashes—no fancy characters or spaces. For example: 'JohnDoe_Resume.docx'.")

    # ----------- PDF HANDLING -----------
    if resume.content_type == 'application/pdf' or resume_name.lower().endswith('.pdf'):
        with pdfplumber.open(resume.file) as pdf:
            fonts_found = set()
            long_paragraphs = 0
            date_format_issues = 0

            for page in pdf.pages:
                for char in page.chars:
                    font_name = char.get('fontname', '').lower()
                    if 'arial' in font_name:
                        fonts_found.add('arial')
                    elif 'calibri' in font_name:
                        fonts_found.add('calibri')
                    elif 'times' in font_name:
                        fonts_found.add('times new roman')
                    else:
                        fonts_found.add(font_name.split('+')[-1])

                text = page.extract_text() or ""
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                for para in paragraphs:
                    if len(para.split()) > 40:
                        long_paragraphs += 1
                    dates = re.findall(r'\b.*?(\d{4}|\d{2}/\d{4}|[A-Za-z]{3,9}\s+\d{4}).*?\b', para)
                    for date in dates:
                        if not any(re.search(pat, date) for pat in DATE_PATTERNS):
                            date_format_issues += 1

            if fonts_found - STANDARD_FONTS:
                deductions += 5
                feedback.append("• Stick to system-standard fonts like Arial, Calibri, or Times New Roman—ATS systems are tuned to read them reliably.")

            if any(page.extract_tables() for page in pdf.pages):
                deductions += 5
                feedback.append("• Reconsider using tables. They often confuse ATS systems and mess with your resume’s readability.")

            if any(page.images for page in pdf.pages):
                deductions += 5
                feedback.append("• Ditch the images—ATS doesn’t see them, and they can interfere with parsing your text.")

            if long_paragraphs:
                feedback.append("• Break down those long paragraphs into bullet points. It’ll boost readability and help recruiters scan faster.")

            if date_format_issues:
                feedback.append("• Stick to one clear format for dates like 'Jan 2020 – Mar 2022'. ATS and human eyes both appreciate the consistency.")

    # ----------- DOCX HANDLING -----------
    elif resume.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or resume_name.lower().endswith('.docx'):
        doc = Document(resume.file)
        resume.file.seek(0)
        fonts_found = set()
        long_paragraphs = 0
        date_format_issues = 0
        header_footer_text_found = False

        def get_effective_font(run, paragraph, doc):
            if run.font.name:
                return run.font.name
            if run.style and hasattr(run.style, "font") and run.style.font.name:
                return run.style.font.name
            if paragraph.style and hasattr(paragraph.style, "font") and paragraph.style.font.name:
                return paragraph.style.font.name
            try:
                normal_style = doc.styles['Normal']
                if hasattr(normal_style, "font") and normal_style.font.name:
                    return normal_style.font.name
            except Exception:
                pass
            return None

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                font = get_effective_font(run, paragraph, doc)
                if font:
                    fonts_found.add(font.strip().lower())
            if len(paragraph.text.split()) > 40:
                long_paragraphs += 1
            dates = re.findall(r'\b.*?(\d{4}|\d{2}/\d{4}|[A-Za-z]{3,9}\s+\d{4}).*?\b', paragraph.text)
            for date in dates:
                if not any(re.search(pat, date) for pat in DATE_PATTERNS):
                    date_format_issues += 1

        if fonts_found - STANDARD_FONTS:
            deductions += 5
            feedback.append("• Stick to system-default fonts—Arial, Calibri, Times New Roman. Anything else risks rendering issues in ATS.")

        if doc.tables:
            deductions += 5
            feedback.append("• Tables might seem neat, but they often mess up parsing. Flatten content into plain text where possible.")

        if hasattr(doc, 'inline_shapes') and doc.inline_shapes:
            deductions += 5
            feedback.append("• Avoid putting in pictures or graphics. ATS won’t read them, and they might block vital info.")

        if long_paragraphs:
            feedback.append("• Break long chunks of text into digestible lines or bullet points—it’s easier on the reader and better for ATS too.")

        if date_format_issues:
            feedback.append("• Unify your date style. Something like 'Jan 2020 – Mar 2022' reads clean and looks polished.")

        try:
            for section in doc.sections:
                header_text = " ".join([p.text for p in section.header.paragraphs]).strip()
                footer_text = " ".join([p.text for p in section.footer.paragraphs]).strip()
                if header_text or footer_text:
                    header_footer_text_found = True
                    break
        except Exception:
            pass

        if header_footer_text_found:
            feedback.append("Keep crucial info out of headers and footers. ATS tools sometimes skip these sections altogether.")

    # ----------- FINAL SCORING -----------
    formatting_points = max(max_points - deductions, 0)
    if not feedback:
        feedback.append("Great job! Your resume formatting looks clean, consistent, and ready for ATS scanning.")

    return (formatting_points, feedback)
